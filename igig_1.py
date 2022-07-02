from docx import Document
import random as rand
import fractions
import math
import tkinter as Tk
import tkinter.filedialog as fd
import re
import tkinter.messagebox as mb
tk=Tk.Tk()
tk.title('Генератор вариантов')

canvas = Tk.Canvas(tk, width = 550, height = 300)
canvas.pack()

a=0
docum = Document()
documAns = Document()

answer = [0 for i in range(13)]#глобальный массив для ответов
#ЗАДАЧА_1
def task1(var):
    if var%2==1:
        isTask = rand.randint(1,2)       
        isPoint = rand.randint(1,2)       
        taskFirst = ["""1) Имеется пять отрезков, длины которых соответственно равны""",
                     """ Наугад берут три из них. Какова вероятность того, что """]
        taskFirstA = ["первый отрезок будет длиной ",
                      ", а второй — "]
        taskFirstB = "из этих отрезков можно построить треугольник?"
        taskSecond = ["1) Среди десяти подарков к Новому году ","""с красной икрой,""",
                      " — с черной и ",
                      """ с икрой заморской, баклажанной. Какова вероятность того, что среди трех наугад взятых подарков """]
        taskSecondA = "два содержат красную икру?"
        taskSecondB = "все три подарка с разной икрой?"        
        if isTask==1:#если выпала первая задача
            cifr11=[]
            for i in range(5):
                r=rand.randint(1,15)
                while r in cifr11:
                    r=rand.randint(1,15)
                cifr11.append(r)
            s=""
            if isPoint==1:#рассматриваем разные пункты
                s+=taskFirstA[0]+str(cifr11[0])+taskFirstA[1]+str(cifr11[1])+" см."
            else:
                s=taskFirstB 
            cifr11.sort()            
            if isPoint==1:#ответы
                answer[0] = "1/20"
            else:
                kl=0
                for i in range(3):
                    for j in range(i+1,4):
                        for k in range(j+1,5):
                            if cifr11[i]+cifr11[j]>cifr11[k]:
                                if cifr11[i]+cifr11[k]>cifr11[j]:
                                    if cifr11[j]+cifr11[k]>cifr11[i]:
                                        kl+=1
                answer[0]=str(kl/10)                                                                   
                
            listCifr = "" #строка с цифрами
            for i in range(4):
                listCifr+=" "+str(cifr11[i])+","
            listCifr += " "+str(cifr11[4])+' см.' 
            ex = taskFirst[0]+listCifr+taskFirst[1]+s
            docum.add_paragraph(ex)            
        elif isTask==2:
            cifr11=[2,2,2]#закинули в минимум по 2 подарка
            r=rand.randint(0,4)
            cifr11[0]+=r#сделали разные значения для разного вида икры
            cifr11[1]+=4-r
            s=""#разбиваем на пункты
            if isPoint==1:
                result = 1
                result*=(cifr11[0]/10)*((cifr11[0]-1)/9)*((cifr11[1]+cifr11[2])/8)
                answer[0]=str(round(result,3))
                s+=taskSecondA
            else:
                s=taskSecondB
                result=1
                result*=(cifr11[0]/10)*(cifr11[1]/9)*(cifr11[2]/8)
                answer[0]=str(round(result,3))                                        
            ex = taskSecond[0]+str(cifr11[0])
            if cifr11[0]>=5:
                ex+=" подарков "
            else:
                ex+=" подарка "
            ex+=taskSecond[1]+str(cifr11[1])+taskSecond[2]+str(cifr11[2])+taskSecond[3]
            ex+=s
            docum.add_paragraph(ex)
    elif var%2==0:
        isTask = rand.randint(1,2)#какое из двух заданий        
        isPoint = rand.randint(1,2)#пункт а или б        
        taskFirst = ["1) К подъезду Транспортной академии в случайном порядке подъезжают "," автомобилей разных марок. Какова вероятноть того, что " ]
        taskFirstA = "первая подъехавшая машина - Таврия, вторая - Мерседес, а третья - Феррари?"
        taskFirstB = "Запорожец подъедет раньше Порше?"
        taskSecond = ["1) На тридцати карточках нарисованы многоугольники: ", " выпуклых, из которых 10 правильных выпуклых и ",
                      """ невыпуклых. Найти вероятность того, что на пяти наугад выбранных карточках окажутся нарисованы """]
        taskSecondA = "три правильных многоугольника?"
        taskSecondB = "два правильных многоугольника и два невыпуклых?"
        if isTask==1:#если выпала первая задача            
            a=rand.randint(5,20)
            s=""
            s+=taskFirst[0]+str(a)+taskFirst[1]            
            if isPoint==1:#рассматриваем разные пункты
                s+=taskFirstA
                ot = fractions.Fraction(1, a*(a-1)*(a-2))                
                answer[0]=str(ot)
            elif isPoint==2:
                s+=taskFirstB
                answer[0]="1/2"
            docum.add_paragraph(s)       
        elif isTask==2:
            a=rand.randint(11,25)#выпуклых всего
            b=a-10#выпуклых неправильных    10-выпуклых правильных
            c=30-a#невыпуклые
            s=""
            s+=taskSecond[0]+str(a)+taskSecond[1]+str(c)+taskSecond[2]   
           
            if isPoint==1:#рассматриваем разные пункты
                s+=taskSecondA
                ot=(math.factorial(10)*math.factorial(20)*math.factorial(5)*math.factorial(25))/(math.factorial(3)*math.factorial(7)*math.factorial(2)*math.factorial(18)*math.factorial(30))
                ot_1=round(ot,2)
                answer[0]=str(ot_1)
            elif isPoint==2:
                s+=taskSecondB
                ot_1=(math.factorial(10)*math.factorial(c)*math.factorial(20-c)*math.factorial(5)*math.factorial(25))/(math.factorial(2)*math.factorial(8)*math.factorial(2)*math.factorial(c-2)*math.factorial(20-c-1)*math.factorial(30))
                ot=round(ot_1,2)
                answer[0]=str(ot)
            docum.add_paragraph(s)           
           
def task2(var):
    if var%2==1: #ВАРИАНТ НЕЧЁТНЫЙ
        isTask = rand.randint(1,2)
        #ПЕРВАЯ ЗАДАЧА
        taskFirst = ["""2) Две фотомодели снимаются для журнала мод «Русская краса», первая — с вероятностью""",""", вторая — с вероятностью """,
                     """. Какова вероятность того, что в январском номере журнала появятся снимки """]
        taskFirstA = 'обеих девушек.'
        taskFirstB = 'только первой девушки.'
        taskFirstC = 'хотя бы одной из них.'
        #ВТОРАЯ ЗАДАЧА
        taskSecond = ["2) Садовод ранней весной высадил саженцы "," яблонь и "," груш. ","Вероятность,что приживется саженец груши, равна ",
                      ", яблони — ",". Какова вероятность, что груш и яблонь приживется поровну?"]
        if isTask == 1: #ЕСЛИ ВЫПАЛА ПЕРВАЯ ЗАДАЧА
            isPoint = rand.randint(1,3)
            a=round(rand.uniform(0.1,1),1)
            b=round(rand.uniform(0.1,1),1)
            ex=taskFirst[0]+str(a)+taskFirst[1]+str(b)+taskFirst[2]
            if isPoint == 1:
                ex+=taskFirstA
                ot=round(a*b,2)
                answer[1]=ot
            elif isPoint == 2:
                ex+=taskFirstB
                notb=1-b
                ot=round(a*notb,2)
                answer[1]=ot
            elif isPoint == 3:
                ex+=taskFirstC
                nota=1-a
                notb=1-b
                ot=round((nota*b)+(a*notb),2)
                answer[1]=ot                                                
            docum.add_paragraph(ex)
            answer[1]=ot
        elif isTask == 2: #ЕСЛИ ВЫПАЛА ВТОРАЯ ЗАДАЧА
            k1=3
            a=round(rand.uniform(0.1,1),1)
            b=round(rand.uniform(0.1,1),1)
            ex=taskSecond[0]+str(k1)+taskSecond[1]+str(k1)+taskSecond[2]+taskSecond[3]+str(a)+taskSecond[4]+str(b)+taskSecond[5]
            #РЕШАЮ ОТВЕТЫ
            p1ya=3*a*(1-a)*(1-a)
            p2ya=3*a*a*(1-a)
            p3ya=a*a*a
            p0ya=(1-a)*(1-a)*(1-a)
            p1g=3*b*(1-b)*(1-b)
            p2g=3*b*b*(1-b)
            p3g=b*b*b
            p0g=(1-b)*(1-b)*(1-b)
            ot=round(p1ya*p1g+p2ya*p2g+p3ya*p3g+p0ya*p0g,3)
            answer[1]=ot                                                   
            docum.add_paragraph(ex)            
    elif var%2==0:#ВАРИАНТ ЧЁТНЫЙ
          isTask=rand.randint(1,2)
          #ПЕРВАЯ ЗАДАЧА
          taskFirst = ["2) Заболевшего студента с одинаковой вероятностью "," могут навестить его друзья и заместитель декана. Какова вероятность того, что в тяжелые для студента дни"]
          taskFirstA = ' его посетит только замдекана?'
          taskFirstB = ' никто не посетит?'
          taskFirstC = ' посетит хотя бы кто-нибудь?'
          #ВТОРАЯ ЗАДАЧА
          task2=['2) Два стрелка делают по два выстрела в мишень. Вероятность попадания в десятку для первого спортсмена равна ',', для второго — ',
                 '. Какова вероятность, что у первого стрелка промахов меньше, чем у второго?']
          if isTask == 1: #ВЫПАЛА ПЕРВАЯ ЗАДАЧА
              isPoint = rand.randint(1,3)
              a=round(rand.uniform(0.1,1),1)
              ex=taskFirst[0]+str(a)+taskFirst[1]
              #РЕШАЮ ОТВЕТЫ
              if isPoint == 1:
                  ex+=taskFirstA
                  ot=round(a*(1-a),2)
                  answer[1]=ot                                           
              elif isPoint == 2:
                  ex+=taskFirstB
                  ot=round((1-a)*(1-a),2)
                  answer[1]=ot                                            
              elif isPoint == 3:
                  ex+=taskFirstC
                  ot=round((1-a)*a*2,2)
                  answer[1]=ot                                            
              docum.add_paragraph(ex)              
          elif isTask == 2: #ВЫПАЛА ВТОРАЯ ЗАДАЧА
                a=round(rand.uniform(0.1,1),1)
                b=round(rand.uniform(0.1,1),1)
                ex=task2[0]+str(a)+task2[1]+str(b)+task2[2]
                #РЕШАЮ ОТВЕТЫ
                ot1=(a*(1-a)*2)*((1-b)*(1-b))
                ot2=(a*a)*((1-b)*(1-b))
                ot3=(a*a)*(b*(1-b)*2)
                ot=round(ot1+ot2+ot3,2)
                answer[1]=ot                                                      
                docum.add_paragraph(ex)                

def task3(var):
    if var%2==1: #НЕЧЕТНЫЙ ВАРИАНТ
        isTask = rand.randint(1,2)#какое из двух заданий        
        #ПЕРВАЯ ЗАДАЧА                      
        taskFirst = ["3) В корзине "," красных и ", """ синих шаров. Из корзины дважды вынимают по одному шару, не кладя их обратно.Найти вероятность появления красного шара при втором испытании, если при первом был извлечен синий шар."""]
        
        taskSecond = ["""3) В эпоху мезолита (среднего каменного века) для того, чтобы убить зайца, было достаточно двух попаданий из лука, при одном попадании вероятность поражения зайца равнялась """,
                      """ Какова вероятность того, что два охотника не остались бы без рагу из зайца, если бы они стреляли по цели из луков одновременно с вероятностью попадания """, " и "," соответственно?"]      
        
        if isTask==1:#если выпала первая задача
            cifr11=[]
            for i in range(2):
                r=rand.randint(2,10)
                while r in cifr11:
                    r=rand.randint(2,10)
                cifr11.append(r)
            s=""
            s+=taskFirst[0]+str(cifr11[0])+taskFirst[1]+str(cifr11[1])+taskFirst[2]
            docum.add_paragraph(s)             
            #РЕШЕНИЕ ЗАДАЧИ                   
            ot = fractions.Fraction(cifr11[0], (cifr11[0]+cifr11[1]-1))
            answer[2]=str(ot)
        elif isTask==2:#если выпала первая задача
            a=round(rand.uniform(0.1,0.9),1)
            b=round(rand.uniform(0.1,0.9),1)
            c=round(rand.uniform(0.1,0.9),1)
            s=""
            s+=taskSecond[0]+str(a)+taskSecond[1]+str(b)+taskSecond[2]+str(c)+taskSecond[3]
            docum.add_paragraph(s) 
            #РЕШЕНИЕ ЗАДАЧИ                   
            ot_1 = a*(1-b)*c+b*(1-a)*c+a*b
            ot=round(ot_1,2)
            answer[2]=str(ot)
    elif var%2==0:
        isTask=rand.randint(1,2)        
          #ПЕРВАЯ ЗАДАЧА
        taskFirst = ["3) В супермаркете на полке лежат ", " плиток белого и ",
                     """ плиток темного шоколада. Покупатель взял, не глядя, сначала одну, затем вторую шоколадку. Найдите вероятность того, что первая из взятых плиток белая, а вторая темная"""]
        taskSecond = ["3) Иван Царевич подъехал к развилке дорог. На камне он прочитал: Налево поехать - женатому быть с вероятностью ", ", прямо - ", ", направо - ",
                      """, а назад уже пути нет. Какова вероятность остаться Ивану Царевичу холостым, если дорогу на развилке он выбрал на удачу?"""]
        if isTask==1:#если выпала первая задача
            cifr11=[]
            for i in range(2):
                r=rand.randint(2,10)
                while r in cifr11:
                    r=rand.randint(2,10)
                cifr11.append(r)
            s=""
            s+=taskFirst[0]+str(cifr11[0])+taskFirst[1]+str(cifr11[1])+taskFirst[2]
            docum.add_paragraph(s)
            #РЕШЕНИЕ ЗАДАЧИ
            ot_1 = (cifr11[0]/(cifr11[0]+cifr11[1]))*(cifr11[1]/(cifr11[0]+cifr11[1]-1))
            ot=round(ot_1,2)
            answer[2]=str(ot)            
        elif isTask==2:#если выпала первая задача
            a=round(rand.uniform(0.1,0.9),1)
            b=round(rand.uniform(0.1,0.9),1)
            c=round(rand.uniform(0.1,0.9),1)
            s=""
            s+=taskSecond[0]+str(a)+taskSecond[1]+str(b)+taskSecond[2]+str(c)+taskSecond[3]
            docum.add_paragraph(s)
            a_1 = 1-a
            b_1 = 1-b
            c_1 = 1-c
            ot_1 =(a_1+b_1+c_1)/3
            ot=round(ot_1,2)
            answer[2]=str(ot)
            
def fac(n):
    if n == 0:
        return 1
    return fac(n-1) * n
 
def task4(var):
    if var%2==1: #ВАРИАНТ НЕЧЁТНЫЙ
        isTask = rand.randint(1,2)
        #ПЕРВАЯ ЗАДАЧА
        if isTask == 1:
            task1 = ['4) По данным ООО «Бытовые услуги», в течение гарантийного срока выходит из строя в среднем ',
                     '% холодильников. Какова вероятность того, что в партии из 100 холодильников не менее половины проработает гарантийный срок?']
            a=rand.randint(1,20)
            p=a/100
            ex=task1[0]+str(a)+task1[1]
            #РЕШЕНИЕ ПО ТЕОРЕМЕ ЛАПЛАСА( P(M1<=M<=M2)=F(X1)-F(X2) m1=50 m2=100
            x1=round(((50-(100*p))/(math.sqrt(100*p*(1-p)))),2)
            x2=round(((100-(100*p))/(math.sqrt(100*p*(1-p)))),2)
            answer[3]="Ф("+str(x1)+')-Ф('+str(x2)+')'
            docum.add_paragraph(ex)            
         #2 ЗАДАЧА
        elif isTask == 2:
            task2 = ['4)  Завод отправил на базу ', ' доброкачественных изделий. Вероятность того, что в пути товар повредится, равна ',
                     '. Найти вероятность того, что на базу поступят ', ' негодных изделия.']
            ysl=rand.randint(2,4)
            n=ysl*10000
            a=round(rand.uniform(0.0001,0.0002),4)
            b=rand.randint(2,4)
            ex=task2[0]+str(n)+task2[1]+str(a)+task2[2]+str(b)+task2[3]
            #РЕШАЮ ПО ПУАССОНУ
            l=n*a
            ot=round(((l** b)/fac(b)*math.exp(l*(-1))),3)
            answer[3]=ot
            docum.add_paragraph(ex)            
    elif var%2==0: #ВАРИАНТ ЧЁТНЫЙ
          isTask = rand.randint(1,2)
          if isTask == 1:
              task1=['4) Среди выпускаемых деталей бывает в среднем ','% брака. Какова вероятность того, что среди взятых на испытание пяти деталей будет ',
                     '% бракованных?']
              a=rand.randint(2,9)
              p=a/100
              c=(rand.randint(1,3)*2)*10
              ex=task1[0]+str(a)+task1[1]+str(c)+task1[2]
              #РЕШЕНИЕ ПО БЕРНУЛЛИ
              m=5*(c/100)
              z=5-m
              c5m=fac(5)/(fac(m)*fac(z))
              ot=round(c5m*(p**m)*((1-p)**z),4)
              answer[3]=ot
              docum.add_paragraph(ex)              
          elif isTask == 2:
              task2=['4) Вероятность для любого абонента позвонить на коммутатор в течение часа равна ','. Телефонная станция обслуживает ',
                    ' абонентов. Какова вероятность того, что в течение часа позвонят ',' абонентa?']
              a=round(rand.uniform(0.0001,0.0002),4)
              ysl=rand.randint(2,4)
              n=ysl*10000
              b=rand.randint(2,4)
              ex=task2[0]+str(a)+task2[1]+str(n)+task2[2]+str(b)+task2[3]
              #РЕШЕНИЕ
              l=n*a
              ot=round(((l** b)/fac(b)*math.exp(l*(-1))),3)
              answer[3]=ot
              docum.add_paragraph(ex)
              
def task5(var):
    if var%2==1: #НЕЧЁТНЫЙ ВАРИАНТ    
        task1 =['5) Охотник, имеющий три патрона, стреляет по дичи до первого попадания или пока не израсходует все патроны.',' Составить ряд распределения числа выстрелов,',
                ' производимых охотником, если вероятность попадания в цель при одном выстреле равна ',
                '. Найти М(Х), D(X), σ(X), этой случайной величины. Построить график F(X).']
        a=round(rand.uniform(0.1,0.9),1)
        b=rand.randint(3,4)
        ex=task1[0]+task1[1]+task1[2]+str(a)+task1[3]
        docum.add_paragraph(ex)
        #ОТВЕТЫ
        x=[0 for i in range(b)]
        z=[0 for i in range(b)]
        for i in range(b):
            x[i]=str(i+1)+' '
            z[i]=i+1
        if b == 3:
            y=[0 for i in range(b)]
            y[0]=a
            y[1]=round((1-a)*a,2)
            y[2]=round(((1-a)*(1-a)*a+(1-a)*(1-a)*(1-a)),2)
        elif b == 4:
            y=[0 for i in range(b)]
            y[0]=a
            y[1]=round((1-a)*a,2)
            y[2]=round((1-a)*(1-a)*a,2)
            y[3]=round(((1-a)*(1-a)*(1-a)*a+(1-a)*(1-a)*(1-a)*(1-a)),2)
        k=0
        for i in range(b):
            k=k+(z[i]*y[i])
        k2=0
        for i in range(b):
            k2=k2+(z[i]*z[i]*y[i])
        d=k2-(k*k)
        ot1='M(x)='+str(round(k,2))+'  D(x)='+str(round(d,2))+'  σ(x)='+str(round(math.sqrt(d),2))
        k=" X= "
        for i in range(len(x)):
            k+=str(x[i])+" "
        k+="""
       Y= """        
        for i in range(len(y)):
            k+=str(y[i])+"  "
        k+="""
       """+ot1
        answer[4]=k        
    elif var%2==0: #ЧЕТНЫЙ ВАРИАНТ   
        task2=['5) В лотерее на ',' билетов разыгрываются четыре вещи, стоимость которых равна 2000, 1000, 500 и 250 руб.',
               'Составить ряд распределения суммы выигрыша для лица, имеющего один билет.',
               ' Найти М(Х), D(X),σ(X) этой случайной величины. Построить график F(X).']
        a=rand.randint(1,4)
        a*=1000
        ex=task2[0]+str(a)+task2[1]+task2[2]+task2[3]
        docum.add_paragraph(ex)
        #ОТВЕТЫ
        x=[0 for i in range(4)]
        z=[0 for i in range(4)]
        x[0]=250
        x[1]=500
        x[2]=1000
        x[3]=2000
        for i in range(4):
            z[i]=' '+str(x[i])+'  '
        y=[0 for i in range(4)]
        m=[0 for i in range(4)]
        for i in range (4):
            y[i]=fractions.Fraction(1,a)
            m[i]=str(fractions.Fraction(1,a))
        k=0
        for i in range(4):
            k=k+(x[i]*y[i])
        k2=0
        for i in range(4):
            k2=k2+(x[i]*x[i]*y[i])
        d=k2-(k*k)
        ot1='M(x)='+str(k)+'  D(x)='+str(k2)+'  σ(x)='+str(round(math.sqrt(d),2))        
        k=" X= "
        for i in range(len(z)):
            k+=str(z[i])+" "
        k+="""
       Y= """        
        for i in range(len(m)):
            k+=str(m[i])+"  "
        k+="""
       """+ot1
        answer[4]=k
        
def task6(var):
    if var%2==1: #НЕЧЁТНЫЙ ВАРИАНТ
        task1=['6) В партии деталей ','% нестандартных. Наугад отобраны четыре детали. ',
               ' Составить ряд распределения случайной величины X — числа нестандартных деталей среди четырех отобранных.',
               ' Найти M(X) и D(X) этой случайной величины.']
        a=rand.randint(1,4)*10
        ex=task1[0]+str(a)+task1[1]+task1[2]+task1[3]
        a/=100
        q=1-a
        docum.add_paragraph(ex)
        #РЕШЕНИЕ
        x=[0 for i in range(5)]
        z=[0 for i in range(5)]
        y=[0 for i in range(5)]
        for i in range(5):
            x[i]=i
            z[i]='  '+str(i)+' '
        c41=fac(4)/fac(3)
        c42=fac(4)/(fac(2)*fac(2))
        y[0]= round(q*q*q*q,4)
        y[1]=round(c41*a*q*q*q,4)
        y[2]=round(c42*a*a*q*q,4)
        y[3]=round(c41*a*a*a*q,4)
        y[4]=round(a*a*a*a,4)
        k=0
        for i in range(5):
            k=k+(x[i]*y[i])
        k2=0
        for i in range(5):
            k2=k2+(x[i]*x[i]*y[i])
        d=k2-(k*k)
        ot1='M(x)='+str(round(k,4))+'  D(x)='+str(round(d,4))        
        k=" X= "
        for i in range(len(z)):
            k+=str(z[i])+" "
        k+="""
       Y= """        
        for i in range(len(y)):
            k+=str(y[i])+"  "
        k+="""
       """+ot1
        answer[5]=k
    elif var%2==0:
        task2=['6) Вероятность того, что вещь, взятая напрокат, будет возвращена исправной, равна ',
               '. Было выдано 4 вещи. Составить ряд распределения числа вещей, которые ',
               'будут возвращены исправными. Найти M(X) и D(X) этой случайной величины.']
        a=round(rand.uniform(0.1,0.9),1)
        q=1-a
        ex=task2[0]+str(a)+task2[1]+task2[2]
        docum.add_paragraph(ex)
        #РЕШЕНИЕ
        x=[0 for i in range(5)]
        z=[0 for i in range(5)]
        y=[0 for i in range(5)]
        for i in range(5):
            x[i]=i
            z[i]='  '+str(i)+' '
        c41=fac(4)/fac(3)
        c42=fac(4)/(fac(2)*fac(2))
        y[0]= round(q*q*q*q,4)
        y[1]=round(c41*a*q*q*q,4)
        y[2]=round(c42*a*a*q*q,4)
        y[3]=round(c41*a*a*a*q,4)
        y[4]=round(a*a*a*a,4)
        k=0
        for i in range(5):
            k=k+(x[i]*y[i])
        k2=0
        for i in range(5):
            k2=k2+(x[i]*x[i]*y[i])
        d=k2-(k*k)
        ot1='M(x)='+str(round(k,4))+'  D(x)='+str(round(d,4))        
        k=" X= "
        for i in range(len(z)):
            k+=str(z[i])+" "
        k+="""
       Y= """        
        for i in range(len(y)):
            k+=str(y[i])+"  "
        k+="""
       """+ot1
        answer[5]=k

def task7(var):
    if var%2==1:
        task1=['7) Вероятность выпуска сверла повышенной хрупкости (брак) равна ',
               '. Сверла укладываются в коробки по 3 штуки. Составить ряд распределения бракованных',
               ' сверл в одной коробке. Найти M(X) этой случайной величины.']
        p=round(rand.uniform(0.01,0.1),2)
        q=1-p
        ex=task1[0]+str(p)+task1[1]+task1[2]
        docum.add_paragraph(ex)
        #РЕШЕНИЕ
        x=[0 for i in range(4)]
        z=[0 for i in range(4)]
        y=[0 for i in range(4)]
        for i in range(4):
            x[i]=i
            z[i]='  '+str(i)+' '
        c31=fac(3)/fac(2)
        y[0]=round(q*q*q,4)
        y[1]=round(c31*p*q*q,4)
        y[2]=round(c31*p*p*q,4)
        y[3]=round(p*p*p,6)
        k=0
        for i in range(4):
            k=k+(x[i]*y[i])
        ot1='M(x)='+str(round(k,4))        
        k=" X= "
        for i in range(len(z)):
            k+=str(z[i])+" "
        k+="""
       Y= """        
        for i in range(len(y)):
            k+=str(y[i])+"  "
        k+="""
       """+ot1
        answer[6]=k        
    elif var%2==0:
        task2=['7) Книга в ',' страниц содержит 3 опечатки. Составить ряд распределения числа опечаток на одной странице.',
               ' Найти M(X) этой случайной величины.']
        a=rand.randint(2,5)
        p=1/4
        q=1-p
        ex=task2[0]+str(a)+task2[1]+task2[2]
        x=[0 for i in range(5)]
        z=[0 for i in range(5)]
        y=[0 for i in range(5)]
        for i in range(5):
            x[i]=i
            z[i]='  '+str(i)+' '
        c51=fac(a)/fac(4)
        y[0]=round(c51*q*q*q*q,4)
        y[1]=round(c51*p*q*q*q,4)
        y[2]=round(c51*p*p*q*q,4)
        y[3]=round(c51*p*p*p*q,4)
        y[4]=round(c51*p*p*p*p,4)
        k=0
        for i in range(5):
            k=k+(x[i]*y[i])
        ot1='M(x)='+str(round(k,4))
        docum.add_paragraph(ex)      

        k=" X= "
        for i in range(len(z)):
            k+=str(z[i])+" "
        k+="""
       Y= """        
        for i in range(len(y)):
            k+=str(y[i])+"  "
        k+="""
       """+ot1
        answer[6]=k

def task8(var):
    docum.add_paragraph("""8) Требуется:
    1) найти плотность вероятности f(x);
    2) построить графики F(x) и f(x);
    3) найти P(α<X<β) для данных α,β. """)    
    if var%2==1:
        docum.add_picture("picture\def8_1.png")
        ans = """            0, x≤2;
       f(x)= 2x/5, 2<x≤3;
                  0, x>3;
       P(2<X<2,5) = F(2,5)-F(2)=0,45"""       
        answer[7]=ans                
    else:
        docum.add_picture("picture\def8_2.png")
        ans = """            0, x≤0;
       f(x) = cos(x), 0<x≤π/2;
                   0, x>ᵴ5/2
       P(0<x<π/6) = F(π/6)-F(0) = 1/2"""        
        answer[7]=ans
        
def task9(var):
    docum.add_paragraph("""9) Требуется:
    1) найти параметр α;
    2) найти функцию распределения F(x);
    3) построить графики F(x) и f(x) """)    
    if var%2==1:
        docum.add_picture("picture\def9_1.png")
        ans ="""a=0,5
                  0, x<2;
      F(x) =x\u00B2-x-2, 2\u2264x\u22643;
                  1, x>3"""
        answer[8]=ans        
    else:
        docum.add_picture("picture\def8_2.png")
        ans = """a=1
                  0, x<0;
      F(x) = (-cos(x)+1)/2, 0\u2264x\u2264\u03C0/2;\n\
                  1, x>\u03C0/2."""                 
        answer[8]=ans

def task10(var):    
    docum.add_paragraph("10) Требуется:")
    docum.add_picture("picture\def10_0.png")
    if var%2==1:
        docum.add_picture("picture\def10_1.png")
        ans = """             0, x\u22643;
         F(x) = 1/8(x\u00B2-6x+9), 3<x\u22645
                      x/2 - 2, 5<x\u22646;
                      1, x>6.
         P(4\u2264x\u22645,5)=5/8=0,625
         M(x) = 59/12=4,92       D(x) = 74/3 - (59/12)^2=0,493       \u03C3(x)=0,702"""
        answer[9]=ans
    else:
        docum.add_picture("picture\def10_2.png")
        ans="""                0, x\u22641;
          F(x)   = 1/3(x\u00B2-2x+1), 1<x\u22642,5;
                        -x\u00B2+6x-8, 2,5<x\u22643;
                         1,x>3
          P(0\u2264x\u22642,5)=3/4
          M(x) = 13/6=2,167      D(x) = 39/8 - (13/6)^2=13/72=0,181       \u03C3(x)=0,425"""
        answer[9]=ans

def task11(var):
    if var%2==1:
        t1=['11) Вероятность выхода из строя гидромуфты валопровода тепловоза за время эксплуатации t задается формулой: ',
            'Случайная величина T — время работы гидромуфты до выхода из строя (в месяцах). Найти среднее время безотказной работы гидромуфты.']
        a=round(rand.uniform(-0.09,-0.01),2)
        ex=t1[0]+'P(t)=1-e^('+str(a)+'t)'+t1[1]
        docum.add_paragraph(ex)
        ot=round(1/(-a),2)
        answer[10]=ot
        
    elif var%2==0:
        t2=['11) Диаметр D детали, изготавливаемой на станке, есть случайная величина, распределенная по нормальному закону (m = ',
            ' см, σ = ',' см). Найти интервал, в котором с вероятностью ',' будут заключены диаметры деталей.']
        a=rand.randint(10,25)
        s=round(rand.uniform(0.1,0.9),1)
        p=round(rand.uniform(0.1,0.9),3)
        ex=t2[0]+str(a)+t2[1]+str(s)+t2[2]+str(p)+t2[3]
        docum.add_paragraph(ex)
        #РЕШЕНИЕ ПО ПРАВИЛУ ТРЕХ СИГМ
        a1=a+(3*s)
        a2=a-(3*s)
        ot=str(a2)+'<X<'+str(a1)
        answer[10]=ot

def task12(var):
    isPoint=rand.randint(1,2)
    if var%2==1:
        p1=['Какова вероятность того, что нагрузка не превысит ',' кг?']
        p2=['Какова вероятность нагрузок от ',' до ',' кг?']
        t1=['12) Нагрузка G на стержень подчиняется нормальному закону распределения с параметрами m = ',
            ' кг; σ = ',' кг. ']
        a=rand.randint(1,5)*50
        b=rand.randint(1,5)*10
        ex=t1[0]+str(a)+t1[1]+str(b)+t1[2]
        if isPoint==1:
            v=rand.randint(20,40)*10
            ex= ex+p1[0]+str(v)+p1[1]
            docum.add_paragraph(ex)
            #ОТВЕТЫ
            x1=round((v-a)/b,2)
            x2=round((0-a)/b,2)
            ot='Ф('+str(x1)+')-Ф('+str(x2)+')'
            answer[11]=ot
        elif isPoint==2:
            v1=rand.randint(1,3)*100
            v2=rand.randint(4,6)*100
            ex+=p2[0]+str(v1)+p2[1]+str(v2)+p2[2]
            docum.add_paragraph(ex)
            #ОТВЕТЫ
            x1=round((v2-a)/b,2)
            x2=round((v1-a)/b,2)
            ot='Ф('+str(x1)+')-Ф('+str(x2)+')'
            answer[11]=ot
    elif var%2==0:
        t2=['12) Время T безотказной работы измерительного комплекса имеет экспоненциальное распределение с математическим ожиданием ',
            ' ч. Какова вероятность того, что комплекс выйдет из строя ']
        p1=['менее чем за ',' ч работы?']
        p2=['не менее чем после ',' ч работы?']
        b=rand.randint(10,20)*100
        ex=t2[0]+str(b)+t2[1]
        if isPoint==1:
            a=rand.randint(1,3)*100
            ex+=p1[0]+str(a)+p1[1]
            docum.add_paragraph(ex)
            #ОТВЕТЫ
            st=-fractions.Fraction(a,b)
            ot='-e^('+str(st)+')+1'
            answer[11]=ot
        elif isPoint==2:
            a=rand.randint(1,3)*100
            ex+=p2[0]+str(a)+p2[1]
            docum.add_paragraph(ex)
            #ОТВЕТЫ
            st=-fractions.Fraction(a,b)
            ot='e^('+str(st)+')'
            answer[11]=ot

def task13(var):
    if var%2==1:
        taskList = ["13) Случайная величина X распределена нормально с математическим ожиданием, равным ",
                    ". Вероятность попадания X в интервал (",
                    ") равна ",
                    ". Чему равна вероятность попадания X в интервал (",
                    ")? Записать для случайной величины X формулу плотности вероятности f(x)."]
        listP=[0.9973, 0.8944]
        xsigm = [2,3]
        a=rand.randint(2,5)
        sigm=rand.randint(2,10)
        isP = rand.randint(0,1)

        b=xsigm[isP]*sigm+a
        alf = 2*a-b
        b2=b+rand.randint(15,25)
        alf2 = b2-rand.randint(5,10)

        s=taskList[0]+str(a)+taskList[1]+str(alf)+";"+str(b)+taskList[2]
        s+=str(listP[isP])+taskList[3]+str(alf2)+";"+str(b2)+taskList[4]
        docum.add_paragraph(s)
        
        
        ans = "P("+str(alf2)+"<x<"+str(b2)+") = \u03A6("+str(round((b2-a)/sigm,2))+")-\u03A6("+str(round((alf2-a)/sigm,2))+")"
        ans+="\n         f(x)=1/("+str(sigm)+"sqrt(2\u03C0))*e^((-x-"+str(a)+")^2/"+str(2*sigm*sigm)+")"+"\n         (a="+str(a)+", \u03C3="+str(sigm)+")"
        answer[12]=ans
    else:
        taskList = ["13) Число вагонов, прибывающих в течение суток на грузовой пункт станции, является случайной величиной, распределенной по нормальному закону с параметрами: a=",                   
                    ",\u03C3=",
                    ". Определить вероятность прибытия на грузовой пункт от ",
                    " до ",
                    " вагонов в сутки."]
        sigm=rand.randint(5,15)
        a=rand.randint(15,40)
        b=rand.randint(10,5*sigm+a)
        alf=rand.randint(-10,b-5)
        s=taskList[0]+str(a)+taskList[1]+str(sigm)+taskList[2]
        s+=str(alf)+taskList[3]+str(b)+taskList[4]
        docum.add_paragraph(s)

        ans="P("+str(alf)+"<x<"+str(b)+") = \u03A6("+str(round((b-a)/sigm,2))+")-\u03A6("+str(round((alf-a)/sigm,2))+")"
        answer[12]=ans

filename = ""
fileDad=""
#filename - фамилии
def fileFam():
    global filename
    filetypes = (("Текстовый файл", "*.docx"), )
    file = fd.askopenfilename(title="Открыть файл", initialdir="/", filetype=filetypes)
    if file:
        filename=file
        lblFam = Tk.Label(text = filename, font=('Times', 11))
        lblFam.place(x=20, y=60, anchor = 'nw')
    else:
        msg = "Выберите файл с фамилиями"
        mb.showerror("Ошибка",msg)
        
def fileGenDir():
    global fileDad
    file = fd.askdirectory(title="Открыть папку", initialdir="/" )
    if file:
        fileDad=file+'/'
#Выбор файла с фамилиями
btn1 = Tk.Button(text = "Выбрать файл с фамилиями", padx = 5,
              pady = 5,background='#B0E0E6',
              activebackground='#E6E6FA', font=('Times', 11),
              command = fileFam)
btn1.place(x=20,y=20, anchor = 'nw')
#учесть вывод ошибки если файл не выбран!!
#Выбор папки для сгенерированного файла
lblDirGenering = Tk.Label(text = "Выберите папку для хранения документа",
                        font=('Times', 12))
lblDirGenering.place(x=250,y=100,anchor = 'nw')
btn2 = Tk.Button(tk,text='Выбрать',padx = 5,
              pady = 3,background='#B0E0E6',
              activebackground='#E6E6FA',
              font=('Times',11), command = fileGenDir)
btn2.place(x=330, y=130,anchor='nw')
#Поле ввода названия файла
name= Tk.StringVar()
lblInputName = Tk.Label(text = 'Введите название файла', font=('Times',13))
lblInputName.place(x=20,y=100,anchor='nw')
txtName=Tk.Entry(textvariable = name)
txtName.place(x=35,y=130,anchor='nw')

#Поле ввода номеров
lblTasks = Tk.Label(text = "Введите номера задач(1-13)", font=('Times',12))
lblTasks.place(x=20,y=180,anchor= 'nw')
number = Tk.StringVar()
txtNumberTasks = Tk.Entry(textvariable = number)
txtNumberTasks.place(x=35,y=205,anchor='nw')

lblAutor = Tk.Label(text = """\tАвторы программы:\n\t Анищенко Екатерина\n  Галаган Яна\n\t    Гончаренко Валентина""",font=('Times', 9))
lblAutor.place(x=370,y=230,anchor = 'nw')


import re
def numTasks(s):
    pattern = r'^[0-9,\-\s]+$'
    is_valid = re.match(pattern,s)
    numberList=[]
    if is_valid:
        sTire = s.split('-')
        if len(sTire)==1:
            s1=sTire[0].split(',')
            for i in range(len(s1)):
                s1[i]=int(s1[i])
            return s1
        for i in range(len(sTire)-1):
            s1=sTire[i].split(',')
            s2=sTire[i+1].split(',')
            for j in range(len(s1)):
                numberList.append(int(s1[j]))
            for j in range(int(s1[len(s1)-1])+1,int(s2[0]),1):
                numberList.append(j)
            if i+1==len(sTire)-1:
                for j in range(len(s2)):
                    numberList.append(int(s2[j]))
        return numberList
    else:
        msg = "Ввод номеров может состоять только из записанных по возрастанию цифр, запятых и тире!"
        mb.showerror("Ошибка",msg)
        return []


#Вывод в документ

def end():
    func = [task1,task2,task3,task4,task5,task6,task7,task8,task9,task10,task11,task12,task13]
    listNumber = numTasks(number.get())
    if listNumber==[]:
        return
    global filename
    if filename=="":
        msg = "Выберите файл с фамилиями"
        mb.showerror("Ошибка",msg)
        return
    document = Document(filename)
    global a
    for paragraph in document.paragraphs:
        a+=1   
        docum.add_paragraph("Вариант №"+str(a)+" - "+paragraph.text)
        for i in range(len(answer)):
            answer[i]=0
        print(listNumber)
        for i in range(len(listNumber)):
            func[listNumber[i]-1](a)
        m=0
        docum.add_page_break()
        documAns.add_paragraph("Ответы на вариант №"+str(a))
        for i in range(len(listNumber)):
            m+=1        
            documAns.add_paragraph(str(listNumber[i])+")  "+str(answer[listNumber[i]-1]))
        documAns.add_page_break()    
    if name.get()=="":
        msg = "Введите название файла"
        mb.showerror("Ошибка",msg)
        return
    fileGen = fileDad+name.get()+'.docx'
    docum.save(fileGen)
    fileAns = fileDad+name.get()+"_ответы.docx"
    documAns.save(fileAns)
    tk.destroy()              
                
btnItog = Tk.Button(text='СГЕНЕРИРОВАТЬ',padx = 5,
              pady = 3,background='#B0E0E6',
              activebackground='#E6E6FA',
              font=('Times',13), command = end)
btnItog.place(x=175, y=230,anchor='nw') 
          
tk.mainloop()
 
             
            


